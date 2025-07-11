from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json
from datetime import datetime
from typing import Literal, Optional, Dict
import os
from docx2pdf import convert

def export_dialog(
    input_json: str,
    speaker_names: Optional[Dict[str, str]] = None,
    file_format: Literal['docx', 'pdf'] = 'docx',
) -> str:
    
    font_name: str = 'Times New Roman'
    speaker_fs: int = 16
    text_fs: int = 14
    time_fs: int = 12
    line_spacing = 1.5
    first_line_indent = Cm(1.25)
    space = Pt(0)

    with open(input_json, 'r', encoding='utf-8') as f:
        segments = json.load(f)  # Изменили название переменной для ясности

    # Группируем слова по speaker и временным отрезкам
    dialog = []
    current_speaker = None
    current_text = []
    
    for segment in segments:
        if 'speaker' not in segment or 'word' not in segment:
            continue  # Пропускаем некорректные сегменты
            
        if segment['speaker'] != current_speaker:
            if current_speaker is not None:
                dialog.append({
                    'speaker': current_speaker,
                    'text': ' '.join(current_text)
                })
            current_speaker = segment['speaker']
            current_text = []
        
        current_text.append(segment['word'])
    
    # Добавляем последнюю реплику
    if current_speaker and current_text:
        dialog.append({
            'speaker': current_speaker,
            'text': ' '.join(current_text)
        })

    doc = Document()
    
    # Настройки стиля по ГОСТ
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(text_fs)
    style.paragraph_format.space_after = space
    style.paragraph_format.space_before = space

    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = line_spacing
    paragraph_format.first_line_indent = first_line_indent
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    date_run = date_para.add_run(datetime.now().strftime("%d.%m.%Y"))
    date_run.font.size = Pt(time_fs)

    for turn in dialog:
        # Обработка имен спикера
        speaker = turn['speaker']
        if speaker_names and speaker in speaker_names:
            speaker = speaker_names[speaker]
        
        # Спикер
        p_speaker = doc.add_paragraph()
        p_speaker.paragraph_format.space_before = space
        p_speaker.paragraph_format.space_after = space
        p_speaker.paragraph_format.first_line_indent = first_line_indent

        speaker_run = p_speaker.add_run(speaker)
        speaker_run.bold = True
        speaker_run.font.size = Pt(speaker_fs)

        # Текст
        p_text = doc.add_paragraph()
        p_text.paragraph_format.line_spacing = line_spacing
        p_text.paragraph_format.space_before = space
        p_text.paragraph_format.space_after = space
        p_text.add_run("— ").bold = True  
        text_run = p_text.add_run(turn['text'])
        text_run.font.size = Pt(text_fs)

    base_name = os.path.splitext(input_json)[0]
    docx_path = f"{base_name}.docx"
    
    doc.save(docx_path)
    
    if file_format == 'pdf':
        pdf_path = f"{base_name}.pdf"
        convert(docx_path, pdf_path)
        return pdf_path
    
    return docx_path